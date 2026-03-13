"""Convert a markdown investment memo to a formatted DOCX file."""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    print(
        "Error: python-docx is required. Install it with: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)


def parse_markdown(text: str) -> list[dict]:
    """Parse markdown into a list of block elements."""
    blocks: list[dict] = []
    lines = text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # Headings
        if m := re.match(r"^(#{1,4})\s+(.+)$", line):
            level = len(m.group(1))
            blocks.append({"type": "heading", "level": level, "text": m.group(2)})
            i += 1
            continue

        # Table (starts with |)
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            blocks.append({"type": "table", "lines": table_lines})
            continue

        # Unordered list item
        if m := re.match(r"^(\s*)[-*]\s+(.+)$", line):
            indent = len(m.group(1))
            level = indent // 2
            blocks.append({"type": "list_item", "text": m.group(2), "level": level})
            i += 1
            continue

        # Ordered list item
        if m := re.match(r"^(\s*)\d+\.\s+(.+)$", line):
            indent = len(m.group(1))
            level = indent // 2
            blocks.append(
                {"type": "ordered_list_item", "text": m.group(2), "level": level}
            )
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+$", line.strip()):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # Code block (skip, not typical in memos)
        if line.strip().startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1
            continue

        # Paragraph (non-empty, non-special line)
        if line.strip():
            para_lines = [line.strip()]
            i += 1
            while (
                i < len(lines)
                and lines[i].strip()
                and not lines[i].strip().startswith("#")
                and not lines[i].strip().startswith("|")
                and not re.match(r"^\s*[-*]\s+", lines[i])
                and not re.match(r"^\s*\d+\.\s+", lines[i])
                and not lines[i].strip().startswith("```")
                and not re.match(r"^---+$", lines[i].strip())
            ):
                para_lines.append(lines[i].strip())
                i += 1
            blocks.append({"type": "paragraph", "text": " ".join(para_lines)})
            continue

        i += 1

    return blocks


def add_formatted_text(paragraph, text: str):
    """Add text with inline markdown formatting (bold, italic) to a paragraph."""
    # Split on bold+italic, bold, and italic markers
    pattern = r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)"
    last_end = 0

    for m in re.finditer(pattern, text):
        # Add text before the match
        if m.start() > last_end:
            paragraph.add_run(text[last_end : m.start()])

        if m.group(2):  # ***bold italic***
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.italic = True
        elif m.group(3):  # **bold**
            run = paragraph.add_run(m.group(3))
            run.bold = True
        elif m.group(4):  # *italic*
            run = paragraph.add_run(m.group(4))
            run.italic = True

        last_end = m.end()

    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def apply_table_style(table):
    """Apply clean formatting to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
            # Header row styling
            if i == 0:
                shading = cell._element
                shading.set(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd",
                    "clear",
                )
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def parse_table_row(line: str) -> list[str]:
    """Parse a markdown table row into cell values."""
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def is_separator_row(line: str) -> bool:
    """Check if a table row is a separator (e.g., | --- | --- |)."""
    cells = parse_table_row(line)
    return all(re.match(r"^[-:]+$", c) for c in cells if c)


def build_docx(blocks: list[dict], output_path: str):
    """Build a DOCX document from parsed blocks."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Configure heading styles
    for level in range(1, 5):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Calibri"
        heading_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        if level == 1:
            heading_style.font.size = Pt(22)
        elif level == 2:
            heading_style.font.size = Pt(16)
        elif level == 3:
            heading_style.font.size = Pt(13)
        else:
            heading_style.font.size = Pt(11)

    for block in blocks:
        match block["type"]:
            case "heading":
                level = min(block["level"], 4)
                doc.add_heading(block["text"], level=level)

            case "paragraph":
                p = doc.add_paragraph()
                add_formatted_text(p, block["text"])

            case "list_item":
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Inches(0.25 * (block["level"] + 1))
                add_formatted_text(p, block["text"])

            case "ordered_list_item":
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.left_indent = Inches(0.25 * (block["level"] + 1))
                add_formatted_text(p, block["text"])

            case "table":
                table_lines = block["lines"]
                # Filter out separator rows
                data_rows = [
                    parse_table_row(line)
                    for line in table_lines
                    if not is_separator_row(line)
                ]
                if not data_rows:
                    continue

                num_cols = max(len(row) for row in data_rows)
                table = doc.add_table(rows=len(data_rows), cols=num_cols)
                table.style = "Light Grid Accent 1"

                for i, row_data in enumerate(data_rows):
                    for j, cell_text in enumerate(row_data):
                        if j < num_cols:
                            cell = table.cell(i, j)
                            cell.text = ""
                            add_formatted_text(cell.paragraphs[0], cell_text)

                apply_table_style(table)
                doc.add_paragraph()  # spacing after table

            case "hr":
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run("─" * 50)
                run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
                run.font.size = Pt(8)

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(
        description="Convert a markdown investment memo to DOCX"
    )
    parser.add_argument(
        "--input",
        required=True,
        help="Path to markdown file",
    )
    parser.add_argument(
        "--output",
        required=True,
        help="Output DOCX file path",
    )
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    markdown_text = input_path.read_text(encoding="utf-8")
    blocks = parse_markdown(markdown_text)
    build_docx(blocks, args.output)
    print(f"Saved: {args.output}")


if __name__ == "__main__":
    main()
